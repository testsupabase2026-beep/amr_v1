# -*- coding: utf-8 -*-
"""
build_namaa_doc.py
==================
Generates the graduation-documentation Word file for the NAMAA Analytics Agent,
following the structure of the teammate's Smart Supermarket ChatBot document.

Run:  python scripts/build_namaa_doc.py
Out:  NAMAA_Analytics_Agent.docx  (project root)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INDIGO   = RGBColor(0x3D, 0x1B, 0x6A)
INDIGO_L = RGBColor(0x8A, 0x45, 0xB2)
CYAN     = RGBColor(0x06, 0xB6, 0xD4)
INK      = RGBColor(0x03, 0x02, 0x13)
MUTED    = RGBColor(0x71, 0x71, 0x82)
GREEN    = RGBColor(0x05, 0x96, 0x69)
CODE_BG  = "F1F0F5"
CALLOUT_BG = "EFE9F6"
BENEFIT_BG = "E9F6EF"

doc = Document()

# ── Base styles ───────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = INDIGO
    st.font.bold = True
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(4)


def _set_page_numbers():
    """Add 'Page X' footer to the default section."""
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    # field: PAGE
    fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
    run._r.append(fld1)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(t): return doc.add_heading(t, level=1)
def h2(t): return doc.add_heading(t, level=2)
def h3(t): return doc.add_heading(t, level=3)
def h4(t): return doc.add_heading(t, level=4)


def _shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _left_border(p, hex_color, size=18):
    """A thick left border → callout/quote bar effect."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), hex_color)
    pbdr.append(left)
    pPr.append(pbdr)


def callout(text, kind="info"):
    """Shaded, left-bar callout box. kind: 'info' (indigo) or 'benefit' (green)."""
    bg = BENEFIT_BG if kind == "benefit" else CALLOUT_BG
    bar = "059669" if kind == "benefit" else "3D1B6A"
    label = "✔ Benefit" if kind == "benefit" else "ℹ Note"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    rb = p.add_run(label + "  ")
    rb.bold = True
    rb.font.size = Pt(10)
    rb.font.color.rgb = GREEN if kind == "benefit" else INDIGO
    r = p.add_run(text)
    r.font.size = Pt(10)
    _shade_para(p, bg)
    _left_border(p, bar)
    return p


def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def cover_page(title, subtitle, meta_lines):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = INDIGO
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle); r2.font.size = Pt(14); r2.font.color.rgb = INDIGO_L
    doc.add_paragraph()
    # accent rule
    rule = doc.add_paragraph(); rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("───────────────────────────────"); rr.font.color.rgb = CYAN
    for _ in range(2):
        doc.add_paragraph()
    for line in meta_lines:
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(line); rl.font.size = Pt(11); rl.font.color.rgb = MUTED
    doc.add_page_break()


def para(t, italic=False, muted=False, size=None, lead=None, bold_prefix=None):
    # `lead`/`bold_prefix`: bold the leading phrase (e.g. "Data source.") then normal text.
    p = doc.add_paragraph()
    _lead = lead or bold_prefix
    if _lead:
        rb = p.add_run(_lead)
        rb.bold = True
        if size:
            rb.font.size = Pt(size)
    r = p.add_run(t)
    r.italic = italic
    if muted:
        r.font.color.rgb = MUTED
    if size:
        r.font.size = Pt(size)
    return p


def bullet(t, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        p.add_run(t)
    else:
        p.add_run(t)
    return p


def numbered(t):
    return doc.add_paragraph(t, style="List Number")


def code(text):
    """Monospace code block with light shading."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "3D1B6A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


_FIG_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "docs", "figures")


def figure_placeholder(label, caption, image=None, width=6.4):
    """Embed a generated figure PNG if available; otherwise leave a labelled placeholder."""
    img_path = os.path.join(_FIG_DIR, image) if image else None
    if img_path and os.path.exists(img_path):
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(img_path, width=Inches(width))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ {label} — insert screenshot here ]")
        r.bold = True
        r.font.color.rgb = MUTED
        r.font.size = Pt(10)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(caption)
    rc.italic = True
    rc.font.size = Pt(9)
    rc.font.color.rgb = MUTED


# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════
cover_page(
    "NAMAA Analytics Agent",
    "Analytics Chatbot — Deep Agent  ·  Text-to-SQL over a Live Data Warehouse",
    [
        "Graduation Project Documentation — Section 4.4",
        "Bilingual (Arabic / English) · Agentic RAG · PostgreSQL DWH · Groq LLMs",
        "",
    ],
)
_set_page_numbers()

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
h1("4.4  Analytics Chatbot — Deep Agent (NAMAA Analytics Agent)")
para(
    "The NAMAA Analytics Agent is an AI-powered, fully agentic analytics assistant that turns "
    "natural-language business questions — in Arabic or English — into validated SQL executed "
    "against a live PostgreSQL data warehouse, and returns a natural-language answer, an "
    "interactive chart, a result table, business recommendations, and suggested follow-up "
    "questions. It removes the need for dashboards or manual SQL: the user simply asks, and the "
    "agent understands, retrieves, computes, visualises, and explains."
)
para(
    "This section documents the agent end to end: what it is and why it exists (4.4.1–4.4.3), "
    "how it is built (4.4.4), a deep walkthrough of each of the eight pipeline stages and the "
    "benefit each one adds (4.4.5), the Text-to-SQL engine at its core (4.4.6), the prompt "
    "engineering that makes the LLM calls reliable (4.4.7), the technology choices (4.4.8), the "
    "front-end and REST back-end (4.4.9), testing and validation (4.4.10), benefits (4.4.11), "
    "platform integration (4.4.12), and the engineering challenges solved along the way (4.4.13)."
)

# ── Architecture gallery ──────────────────────────────────────
h2("System Architecture at a Glance")
para(
    "Three high-level views precede the detailed feature sections: (A) the full-stack system "
    "overview, (B) the internal pipeline (the 8-stage agent flow), and (C) the Text-to-SQL data "
    "flow for a single query."
)
figure_placeholder(
    "Figure A — System Architecture Overview",
    "Figure A. End-to-end architecture: browser UI (or FastAPI REST client) → pipeline "
    "orchestrator → Groq LLMs (RotatingKeyLLM) + FAISS schema index + live PostgreSQL DWH "
    "(Supabase, schema dwh1) → answer + chart + recommendations.",
    image="figure_a_architecture_gen.png",   # Gemini version (accurate labels + icons)
)
figure_placeholder(
    "Figure B — Agent Pipeline (8 stages)",
    "Figure B. The pipeline: chart-edit intercept → spelling/reference/chitchat → exact cache → "
    "rewrite + combined intent-&-decompose → semantic cache → SQL phase "
    "(schema → decompose → SQL → combine) → parallel enrichment (summary + chart + recos + "
    "follow-ups) → finalize (cache store, metrics, final answer).",
    image="figure_b_pipeline.png",
)
figure_placeholder(
    "Figure C — Text-to-SQL Data Flow",
    "Figure C. A single analytics query: FAISS schema retrieval → LLM SQL generation → "
    "multi-layer SQL safety validation → execution on dwh1 → sanity check / self-correction "
    "retry → deterministic chart spec → rendered chart.",
    image="figure_c_text_to_sql_gen.png",   # Gemini version (accurate + clearer retry loop)
)

# ══════════════════════════════════════════════════════════════
# 4.4.1 System Overview
# ══════════════════════════════════════════════════════════════
h2("4.4.1  System Overview")
para(
    "The Analytics Agent is a stage-orchestrated, multi-layer system that converts a "
    "natural-language question into an executed SQL analysis and a fully-formatted answer. Each "
    "user turn passes through an 8-stage asynchronous pipeline: (0) a chart-edit intercept that "
    "modifies the current chart without re-running SQL; (1) spelling correction, pronoun / "
    "follow-up reference resolution, and a chitchat gate; (2) an exact-match cache lookup; (3) a "
    "combined query-rewrite + intent-classification + decomposition call; (4) a semantic "
    "(embedding-similarity) cache lookup; (5) the SQL "
    "phase — FAISS schema retrieval, simple/compound execution, and result combination; (6) a "
    "parallel enrichment stage that generates the NL summary, chart, recommendations, and "
    "follow-up questions concurrently; and (7) finalization — caching, metrics, and the final "
    "answer."
)
para(
    "The agent is bilingual (Arabic + English), multi-user safe (each browser tab / API session "
    "is isolated through a ContextVar-backed session state), and observable (one JSONL metrics "
    "record per query). It is delivered through two front-ends that share the identical pipeline: "
    "a Gradio web app and a framework-free FastAPI REST endpoint with a static HTML UI."
)

# ══════════════════════════════════════════════════════════════
# 4.4.2 Purpose
# ══════════════════════════════════════════════════════════════
h2("4.4.2  Purpose")
para(
    "The agent gives non-technical business users direct, conversational access to a retail data "
    "warehouse. It addresses a concrete problem: meaningful analytics normally require either "
    "pre-built dashboards (which only answer questions someone anticipated) or an analyst who can "
    "write SQL (a bottleneck and a cost). NAMAA removes both. A store manager can ask "
    "“أكثر 5 فئات إيرادات في 2025” or “compare TIDE vs ARIEL revenue in 2024” and receive, in "
    "seconds, an accurate answer computed live from the warehouse — with a chart, a plain-language "
    "explanation, and next-step suggestions."
)
para("Concretely, the agent is designed to:")
bullet("Understand intent and correct typos in either language before doing anything else.")
bullet("Resolve pronouns and short follow-ups from conversation history (e.g. “now only 2024”).")
bullet("Decompose compound questions (e.g. growth, ratio, top-N + trend) into sub-queries.")
bullet("Generate, validate, and execute read-only SQL autonomously against the live DWH.")
bullet("Visualise the result deterministically and explain it, with actionable recommendations.")
bullet("Stay safe (read-only SQL, sandboxed chart code) and bounded (cache, key rotation, retries).")

# ══════════════════════════════════════════════════════════════
# 4.4.3 Functionality
# ══════════════════════════════════════════════════════════════
h2("4.4.3  Functionality")
para(
    "The agent processes a natural-language question (plus session history and three feature "
    "toggles — chart, recommendations, cache) and streams back a fully-formed analytical answer. "
    "The primary REST entry point is POST /ask (blocking) or POST /ask/stream (token-by-token "
    "Server-Sent Events). Every response carries the natural-language answer, an HTML result "
    "table, the chart in three interchangeable forms, business recommendations, three follow-up "
    "questions, and the token count for the call."
)

para("Response fields (final chunk):")
table(
    ["Field", "Type", "Description"],
    [
        ["summary / chat_text", "string", "2–3 sentence natural-language answer (K/M-formatted numbers)"],
        ["result_html", "string", "Styled HTML result table (drop into a <div>)"],
        ["chart_json", "string", "Plotly JSON spec — the primary field for a custom frontend"],
        ["chart_html", "string", "Self-contained Plotly HTML (for an <iframe srcdoc>)"],
        ["reco_text", "string", "Business recommendations (numbered, same language as question)"],
        ["followup", "list[str]", "Three suggested follow-up questions"],
        ["tokens_used", "int", "LLM tokens consumed by this call"],
        ["done", "bool", "True on the final chunk only"],
    ],
    widths=[1.7, 1.0, 4.0],
)

para("Request — Arabic ranking question:")
code(
    '{\n'
    '  "question": "أكثر 5 فئات إيرادات في 2025",\n'
    '  "use_viz": true,\n'
    '  "use_reco": true,\n'
    '  "use_cache": true,\n'
    '  "session_id": "user-123"\n'
    '}'
)

para("Response — ranking question (abridged, real production shape):")
code(
    '{\n'
    '  "summary": "حققت فئة المواد الغذائية أعلى إيرادات في 2025 بنحو 12.58 مليون\\n'
    '              دينار كويتي، تليها المنظفات ثم العناية الشخصية…",\n'
    '  "result_html": "<table class=\'namaa-table\'>…</table>",\n'
    '  "chart_json": "{\\"data\\": [{\\"type\\": \\"bar\\", \\"orientation\\": \\"h\\", …}],\n'
    '                 \\"layout\\": {…}}",\n'
    '  "reco_text": "1. زيادة المخزون للمواد الغذائية …\\n2. مراجعة تسعير …",\n'
    '  "followup": ["ما اتجاه مبيعات المواد الغذائية شهرياً؟", "…", "…"],\n'
    '  "tokens_used": 8250,\n'
    '  "done": true,\n'
    '  "session_id": "user-123"\n'
    '}'
)
para(
    "Note on language. The agent detects the language of the question and produces the summary, "
    "recommendations, and follow-ups in that same language — an Arabic question yields an Arabic "
    "answer, an English question an English one — while SQL, column names, and the chart data "
    "remain language-neutral.", italic=True, size=10,
)

# ══════════════════════════════════════════════════════════════
# 4.4.4 Technical Implementation
# ══════════════════════════════════════════════════════════════
h2("4.4.4  Technical Implementation")
para("The system is a modular Python application. Its main components:")
bullet("Slim async orchestrator that delegates each stage to a dedicated module.",
       bold_prefix="Pipeline (src/pipeline.py): ")
bullet("one file per stage — state, chart_edit, cache_serve, sql_phase, enrichment.",
       bold_prefix="Pipeline stages (src/pipeline_stages/): ")
bullet("PipelineState dataclass (per-request state + timing + yield_dict helper).",
       bold_prefix="State (state.py): ")
bullet("RotatingKeyLLM — 3-key rotation on daily limits + 4-model fallback chain + "
       "transient-error retry; plus a per-session token tracker.",
       bold_prefix="LLM layer (src/llm.py): ")
bullet("spelling, reference resolution, chitchat gate, query rewrite, combined "
       "intent-&-decompose, and the chart-edit gate — with cheap regex pre-checks that skip "
       "the LLM on already-clean input.",
       bold_prefix="Intent layer (src/intent.py): ")
bullet("SQL validation + execution, six hardcoded pandas combiners, top-N enforcement, "
       "and the LLM-code sandbox.",
       bold_prefix="Executor (src/executor.py): ")
bullet("conversation memory, exact + semantic cache, ContextVar multi-user isolation, "
       "and per-session save/load.",
       bold_prefix="Session (src/session.py): ")
bullet("all LangChain prompt templates (SQL, SQL-fix, rewrite, intent-decompose, combine, "
       "recommendations, follow-ups, summary, chart, chart-edit, chitchat, spelling).",
       bold_prefix="Prompts (src/prompts.py): ")
bullet("DWH engine, tuning constants, lazily-loaded FAISS index + multilingual embeddings, "
       "and the NAMAA brand chart theme.",
       bold_prefix="Config (src/config.py): ")
bullet("PDF report, NL summary (streamed), recommendations, follow-up generation.",
       bold_prefix="Export (src/export.py): ")

para("Unlike a CSV-based prototype, every query runs against the live "
     "PostgreSQL DWH (Supabase, schema dwh1) — a star schema of one fact table and five "
     "dimension tables. Only the compact schema description is retrieved (via FAISS) and given "
     "to the LLM; the data itself is never loaded into memory.", lead="Data source. ")

# ══════════════════════════════════════════════════════════════
# 4.4.5 Agent Graph Nodes / Pipeline Stages
# ══════════════════════════════════════════════════════════════
h2("4.4.5  The 8-Stage Pipeline — A to Z")
para(
    "Where a LangGraph-style agent is a graph of nodes, the NAMAA agent is an ordered async "
    "pipeline of eight stages, each a small module (in src/pipeline_stages/) with exactly one "
    "responsibility. The orchestrator (src/pipeline.py) runs them in sequence and lets any stage "
    "short-circuit: a chart-edit, a cache hit, or a chitchat message returns immediately, so the "
    "full, expensive path runs only for a genuine, cache-missing analytics question. This is the "
    "single most important performance property of the design — the common cases are cheap."
)
para(
    "The table below is the map; the sub-sections that follow walk each stage in depth — what it "
    "does, how it works internally, and the concrete benefit it adds to the product.")
table(
    ["#", "Stage", "One-line role", "Can exit early?"],
    [
        ["0", "Chart-edit intercept", "Edit the current chart without new SQL", "Yes"],
        ["1", "Preprocess (spelling · refs · chitchat)", "Clean the question and gate non-analytics", "Yes (chitchat)"],
        ["2", "Exact cache", "Instant replay of an identical question", "Yes (hit)"],
        ["3", "Intent + decompose", "Understand the question in one LLM call", "No"],
        ["4", "Semantic cache", "Replay a paraphrase of a prior question", "Yes (hit)"],
        ["5", "SQL phase", "Retrieve schema, generate & run SQL", "No"],
        ["6", "Parallel enrichment", "Summarise, chart, recommend, suggest", "No"],
        ["7", "Finalize", "Cache, accumulate report, emit metrics", "No"],
    ],
    widths=[0.35, 2.4, 2.6, 1.2],
)

# ── Stage 0 ───────────────────────────────────────────────────
h3("Stage 0 · Chart-Edit Intercept")
para("What it does. Before anything else, the agent checks whether the message is a request to "
     "modify the chart already on screen — “flip to bar”, “make it a pie”, “use green”, "
     "“اقلب لأعمدة” — rather than a new data question.", lead="")
para("How it works. A lightweight gate LLM call (CHART_EDIT_GATE_PROMPT) classifies the message "
     "as chart_edit vs new_query, but only if a previous chart exists in the session "
     "(_last_plotly_code). On a chart-edit, the existing chart code is surgically modified "
     "(CHART_EDIT_PROMPT), re-executed in the sandbox, re-rendered (including a fresh PNG for the "
     "PDF), and returned — the SQL layer is never touched. Any failure falls through to the "
     "normal pipeline.", lead="")
callout("A follow-up like “now make it a pie chart” costs one tiny classification call and a "
        "re-render instead of a full re-query (schema retrieval + SQL generation + execution + "
        "enrichment) — turning an interactive chart tweak from ~30 seconds into ~1 second.",
        kind="benefit")

# ── Stage 1 ───────────────────────────────────────────────────
h3("Stage 1 · Preprocess — Spelling, Reference Resolution, Chitchat Gate")
para("What it does. It cleans the raw question and decides whether it is actually an analytics "
     "request at all. Three sub-steps run in order: (a) spelling/typo correction, (b) pronoun "
     "and follow-up reference resolution using conversation history, (c) a chitchat gate that "
     "diverts greetings and off-topic messages to a friendly reply.", lead="")
para("How it works. Two cheap regex pre-checks make this stage almost free on clean input: "
     "_looks_clean() skips the spelling LLM unless the text has over-long tokens, 3+ repeated "
     "characters, or unexpected symbols; _needs_resolution() skips the reference resolver unless "
     "the text contains a pronoun/demonstrative or a follow-up prefix (“and…”, “now…”, "
     "“والآن…”). Only when a pre-check trips does an LLM call actually fire. The chitchat gate "
     "(CHITCHAT_GATE_PROMPT) then labels the message analytics vs chitchat; a chitchat message is "
     "answered by CHITCHAT_RESPONSE_PROMPT and the pipeline returns.", lead="")
callout("Reference resolution is what makes a real conversation possible: after “top 10 products”, "
        "a bare “now only in 2024” is expanded into a full, self-contained question before it "
        "ever reaches SQL. The regex pre-checks mean this intelligence adds roughly zero latency "
        "to already-clean questions — the LLM is called only when it can actually help.",
        kind="benefit")

# ── Stage 2 ───────────────────────────────────────────────────
h3("Stage 2 · Exact Cache")
para("What it does. It replays the stored answer for a question that is byte-for-byte identical "
     "(after rewriting) to one already answered in the session.", lead="")
para("How it works. A SHA-256 key is computed over the rewritten question plus the "
     "chart/recommendation flags. If that key is present in the session's query cache, the full "
     "cached payload — summary, table, chart, recommendations, follow-ups — is returned "
     "immediately with zero LLM tokens, and the query is still recorded so it appears in the "
     "session's PDF report.", lead="")
callout("Repeated or refreshed questions (very common in real analytics sessions) return in "
        "milliseconds at zero token cost — directly protecting the free-tier Groq budget and "
        "making the UI feel instant.", kind="benefit")

# ── Stage 3 ───────────────────────────────────────────────────
h3("Stage 3 · Intent Classification + Decomposition")
para("What it does. In a single LLM call it both understands the question and plans how to "
     "answer it. It detects the language, rewrites the question into a precise English intent, "
     "and returns a structured JSON describing the query.", lead="")
para("How it works. INTENT_DECOMPOSE_PROMPT returns one JSON object with: intent_type "
     "(ranking / trend / distribution / comparison / correlation / detail), chart_type, "
     "needs_chart, top_n, time_filter, dimension, metric, and — if the question is compound — "
     "is_compound, the sub-step descriptions, and the combination strategy. Combining "
     "classification and decomposition into one call (rather than two) roughly halves the "
     "pre-SQL latency. The output is strictly validated: unknown intents/metrics fall back to "
     "safe defaults, and a malformed plan is repaired before it reaches the SQL stage.", lead="")
callout("This one structured call is the agent's “brain”: every downstream decision — which "
        "cache key to check, which SQL to write, which chart to draw, which language to answer "
        "in — is driven by these fields. Extracting them all at once keeps the agent both fast "
        "and consistent.", kind="benefit")

# ── Stage 4 ───────────────────────────────────────────────────
h3("Stage 4 · Semantic Cache")
para("What it does. It replays a prior answer for a question that is a paraphrase — not "
     "identical text, but the same analytical request (possibly in the other language).", lead="")
para("How it works. The rewritten question is embedded with a multilingual MiniLM model and "
     "compared by cosine similarity against previously-answered questions. A hit requires the "
     "similarity to clear the 0.70 threshold AND an exact match on the structured facets "
     "(top_n, time_filter, dimension, metric, intent_type) plus a period signature — so “top 5 "
     "products 2024” never collides with “top 10 products 2025”. On a cross-language hit the SQL "
     "result and chart are reused directly and only the natural-language text (summary, "
     "recommendations, follow-ups) is regenerated in the new language.", lead="")
callout("An Arabic user and an English user asking the same thing share one cached computation — "
        "the expensive SQL + chart work happens once, and only the cheap prose is re-written per "
        "language. Facet-gating prevents the classic semantic-cache failure of returning a "
        "similar-looking but wrong answer.", kind="benefit")

# ── Stage 5 ───────────────────────────────────────────────────
h3("Stage 5 · SQL Phase (Text-to-SQL)")
para("What it does. This is the analytical core: it retrieves the relevant schema, generates "
     "SQL, validates it for safety, executes it against the live warehouse, and self-corrects on "
     "failure. Simple questions run one query; compound questions run parallel sub-steps that are "
     "merged deterministically.", lead="")
para("How it works. Detailed in Section 4.4.6. In brief: FAISS retrieves the top-K schema chunks; "
     "the LLM writes PostgreSQL; a five-layer validator guarantees a single read-only SELECT with "
     "a row cap; the query executes in a read-only transaction; a sanity check catches "
     "silent-wrong-answer patterns and feeds errors back for up to three retries. Compound "
     "questions decompose into (at most) two symmetric sub-steps executed in parallel and merged "
     "by one of six hardcoded pandas combiners.", lead="")
callout("The user gets an exact, live answer to a question nobody pre-built a dashboard for — "
        "safely (the warehouse can only ever be read), and correctly (validation + self-"
        "correction catch both unsafe and silently-wrong SQL before the user sees it).",
        kind="benefit")

# ── Stage 6 ───────────────────────────────────────────────────
h3("Stage 6 · Parallel Enrichment")
para("What it does. It turns the raw result table into a complete answer: a natural-language "
     "summary, a chart, business recommendations, and three follow-up questions.", lead="")
para("How it works. The four tasks are launched as concurrent asyncio tasks. The NL summary "
     "streams to the user token-by-token while the chart, recommendations, and follow-ups "
     "complete in the background — so wall-clock time is the slowest single task (~4 s), not "
     "their sum (~12 s). The chart is built deterministically in Python from the result's shape "
     "(a chart spec chosen by the pipeline), which is why charts are reliable and never depend on "
     "the LLM writing correct plotting code. All numbers are formatted for readability (K/M), and "
     "the answer is produced in the question's language.", lead="")
callout("The user perceives a fast, streaming answer and receives four value-adds at once — "
        "explanation, visualisation, next actions, and next questions — for the price (in time) "
        "of just the slowest one. Deterministic charting removes an entire class of "
        "LLM-generated-code failures.", kind="benefit")

# ── Stage 7 ───────────────────────────────────────────────────
h3("Stage 7 · Finalize")
para("What it does. It persists everything the session and the platform need: it stores the "
     "cache entry (so Stages 2 and 4 can replay it later), appends this query's recommendation "
     "to the session report, and emits one metrics record.", lead="")
para("How it works. The full payload (result, chart, summary, recommendations, follow-ups, "
     "language, chart PNG path) is written to the query cache and the semantic index; the "
     "recommendation is accumulated for the PDF report; and a JSONL metrics line — question, "
     "language, intent, compound flag, cache outcome, tokens, per-stage timings, total latency, "
     "success — is appended to metrics.jsonl. Metrics writing is best-effort and never breaks "
     "the response.", lead="")
callout("Every answer makes the next one cheaper (caching), contributes to an exportable report "
        "(recommendations), and is fully observable (metrics) — so the system is measurable and "
        "auditable in production, not a black box.", kind="benefit")

spacer()
para(
    "Each request installs its own SessionState into a ContextVar; asyncio propagates that "
    "context to every child task and thread, so concurrent users never share conversation "
    "history, cache, or token counters. A module-level attribute proxy lets existing code read "
    "and write session fields transparently, keeping the multi-user machinery invisible to the "
    "rest of the codebase.", lead="Cross-cutting — multi-user isolation. ")

# ══════════════════════════════════════════════════════════════
# 4.4.6 Text-to-SQL Layer
# ══════════════════════════════════════════════════════════════
h2("4.4.6  Text-to-SQL Layer")
para(
    "The Text-to-SQL layer is the core of the agent. It converts a rewritten, classified "
    "question into a validated, executed SQL statement against dwh1, with autonomous "
    "self-correction. The flow for a simple query is: retrieve schema context → generate SQL → "
    "validate (multi-layer safety) → execute → sanity-check the result → on failure, feed the "
    "error back to the LLM and retry (up to 3 attempts)."
)

h3("Schema retrieval (FAISS)")
para(
    "Rather than dumping the entire schema into every prompt, the top-K most relevant schema "
    "chunks are retrieved from a FAISS index of the dwh1 tables using a multilingual MiniLM "
    "embedding, then injected into the SQL-generation prompt. The FAISS index and the embedding "
    "model are loaded lazily on first use — so the process starts fast and only the request path "
    "pays the load cost."
)

h3("Multi-layer SQL safety")
para("Generated SQL is never blindly executed. Five independent guards protect the warehouse:")
numbered("DDL/DML block — a regex rejects INSERT, UPDATE, DELETE, DROP, ALTER, TRUNCATE, "
         "GRANT, REVOKE, CREATE, COPY, MERGE, CALL, DO, VACUUM, ANALYZE.")
numbered("Single-statement only — semicolon-separated multiple statements are rejected.")
numbered("SELECT / WITH root only — any non-read statement is refused.")
numbered("Row cap — LIMIT 10000 is enforced; larger user limits are capped.")
numbered("Read-only transaction + statement timeout at the connection level — defense-in-depth "
         "if the regex ever misses something.")

h3("Sanity check + self-correction")
para(
    "After execution the result is sanity-checked for silent wrong-answer bugs — an all-zero "
    "numeric result (likely a missing JOIN or wrong filter), a ranking that returned a single "
    "row when several were asked for, or a trend result with no time column. A failed check "
    "feeds the specific error back into the SQL-fix prompt and the query is regenerated, up to "
    "the retry limit. An empty sub-step result also triggers one retry with the emptiness fed "
    "back."
)

h3("Compound decomposition and hardcoded combiners")
para(
    "Compound questions are split into (at most) two symmetric sub-steps that run in parallel, "
    "then merged by one of six deterministic pandas combiners — no LLM call for the merge, which "
    "is faster and eliminates a class of code-generation bugs. The LLM exec() path remains only "
    "as a fallback for genuinely novel shapes."
)
table(
    ["Combination", "What it computes"],
    [
        ["merge_on_key", "Join two periods/segments on a shared key → grouped comparison (also scalar-vs-scalar → labelled 2-bar)."],
        ["pct_change", "% change between two periods; a near-zero base is flagged as a NEW ENTRANT (NaN, not a bogus 180,000%)."],
        ["subtract", "Difference between two periods."],
        ["ratio", "step1 as a percentage of step2 = (part / whole) × 100 (share-of-total)."],
        ["filter_by_step1", "step1 = the top-N entity set; step2 = a detail/trend over all entities; the combiner intersects them by value overlap."],
        ["display_separately", "Two independent tables rendered side-by-side with a sub-plot chart."],
    ],
    widths=[1.6, 5.0],
)
para("All six combiners are implemented as plain pandas (roughly 30 lines each). The dispatcher "
     "tries the hardcoded function first and falls back to the LLM only for genuinely novel "
     "shapes — saving 2–3 seconds and all combine-step tokens on a typical compound query.",
     italic=True, size=10)

h3("Example — generated SQL")
para("For “أكثر 5 فئات إيرادات في 2025” (top 5 categories by revenue in 2025) the agent "
     "generates and executes:")
code(
    "SELECT c.category_name,\n"
    "       ROUND(SUM(f.total_amount)::numeric, 2) AS revenue_kwd\n"
    "FROM dwh1.fact_order_item f\n"
    "JOIN dwh1.dim_category c ON f.category_key = c.category_key\n"
    "JOIN dwh1.dim_date     d ON f.order_date_key = d.date_key\n"
    "WHERE d.year = 2025\n"
    "GROUP BY c.category_name\n"
    "ORDER BY revenue_kwd DESC\n"
    "LIMIT 5;"
)

# ══════════════════════════════════════════════════════════════
# 4.4.7 Prompt Engineering
# ══════════════════════════════════════════════════════════════
h2("4.4.7  Prompt Engineering")
para(
    "Every LLM call in the agent is driven by a dedicated template in src/prompts.py, and each "
    "template is engineered to make the model’s output predictable enough to branch on in code. "
    "The agent does not rely on a single mega-prompt; it uses a small library of single-purpose "
    "prompts, each with one job and one output contract."
)
para("Five design rules are applied consistently across the prompt library:")
bullet("each prompt does exactly one thing (classify, rewrite, generate SQL, summarise), so its "
       "output is easy to validate and reuse.", bold_prefix="Single responsibility. ")
bullet("classification, decomposition, and follow-up prompts must return valid JSON with named "
       "keys wrapped in a fenced example — which the code then parses and validates, rejecting "
       "anything malformed.", bold_prefix="Strict output format. ")
bullet("every user-facing prompt is told to reply in the exact language of the question; SQL and "
       "chart data stay language-neutral.", bold_prefix="Language locking. ")
bullet("generation prompts list explicit prohibitions (never invent a product, never guess a "
       "price, never fabricate SQL columns, never transliterate a brand name) — negative "
       "constraints do most of the anti-hallucination work.", bold_prefix="Explicit “do NOT” rules. ")
bullet("SQL-fix, combine-fix, and chart-fix prompts receive the previous failing output plus the "
       "exact error, so a retry is genuinely different from the first attempt, not a re-sample.",
       bold_prefix="Feedback-aware retries. ")

h3("Prompt-to-stage map")
table(
    ["Prompt", "Stage / caller", "Purpose"],
    [
        ["SPELL_CORRECT_PROMPT", "1 · preprocess", "Fix typos without changing meaning, names, or numbers."],
        ["REFERENCE_RESOLVE_PROMPT", "1 · preprocess", "Expand pronouns / short follow-ups into a self-contained question."],
        ["CHITCHAT_GATE_PROMPT", "1 · preprocess", "Classify analytics vs chitchat."],
        ["CHITCHAT_RESPONSE_PROMPT", "1 · preprocess", "Friendly non-analytics reply in the user’s language."],
        ["REWRITE_PROMPT", "3 · intent", "Rewrite to a precise English intent, preserving entity scripts."],
        ["INTENT_DECOMPOSE_PROMPT", "3 · intent", "Combined intent classification + decomposition (one JSON call)."],
        ["SQL_PROMPT / SQL_FIX_PROMPT", "5 · SQL phase", "Generate PostgreSQL / repair it given the error."],
        ["COMBINE_PROMPT / _FIX", "5 · SQL phase", "Fallback pandas code to merge compound sub-steps."],
        ["PLOTLY_PROMPT / _FIX", "6 · enrichment", "Fallback chart code when no deterministic spec applies."],
        ["NL_SUMMARY_PROMPT", "6 · enrichment", "2–3 sentence summary, language-locked, K/M-formatted."],
        ["BUSINESS_RECO_PROMPT", "6 · enrichment", "3–5 actionable recommendations in the question’s language."],
        ["FOLLOWUP_PROMPT", "6 · enrichment", "Exactly three follow-up questions as a JSON array."],
        ["CHART_EDIT_GATE / CHART_EDIT_PROMPT", "0 · chart-edit", "Detect and apply a chart modification."],
    ],
    widths=[2.4, 1.5, 3.0],
)

h3("Worked example — the SQL prompt")
para(
    "SQL_PROMPT is the most safety-critical template. It embeds the full dwh1 star-schema cheat "
    "sheet (tables, columns, join keys) plus hard rules that turn the LLM from a free generator "
    "into a constrained one: use SUM(f.total_amount) for revenue; filter names with ILIKE on the "
    "core Arabic noun (never exact “=”, which misses the definite article “ال”); match brands "
    "against both the Arabic and Latin name columns; exclude the delivery-date sentinel; and "
    "produce exactly one read-only SELECT. When a query fails, SQL_FIX_PROMPT receives the bad "
    "SQL and the database error and returns a corrected statement."
)
callout(
    "Because the schema, the join keys, and the data caveats (discount is always zero, customer/"
    "seller PII is NULL, brands are bilingual) are all encoded in the prompt, the model generates "
    "correct, safe SQL on the first attempt for the large majority of questions — and self-"
    "corrects on the rest.", kind="benefit")

# ══════════════════════════════════════════════════════════════
# 4.4.8 Technology Rationale
# ══════════════════════════════════════════════════════════════
h2("4.4.8  Technology Rationale")
para(
    "Every technology choice was made against the constraints of a real, cost-bounded "
    "deployment: fast inference, zero GPU budget, a live cloud database, and bilingual support."
)
table(
    ["Layer", "Choice", "Why"],
    [
        ["LLM inference", "Groq (LLaMA-3.3-70B + fallbacks)", "Extremely fast inference; free tier; a 4-model fallback chain absorbs per-model outages."],
        ["Orchestration", "Custom async pipeline + LangChain", "Fine-grained control of streaming, caching, and per-stage timing without heavy framework overhead."],
        ["Embeddings", "multilingual MiniLM (CPU)", "Handles Arabic + English schema retrieval and semantic cache with no GPU."],
        ["Vector store", "FAISS (local)", "Fast in-process similarity search; the schema is small and fixed."],
        ["Database", "PostgreSQL DWH (Supabase)", "Real star schema; SQL is the right language for aggregation; read-only pooled connection."],
        ["Charts", "Plotly (deterministic specs)", "Interactive, JSON-serialisable; built in Python from data shape so the LLM cannot break them."],
        ["API / UI", "FastAPI + static HTML / Gradio", "One pipeline, two front-ends; async-native; SSE streaming; framework-free client integration."],
    ],
    widths=[1.4, 2.0, 3.2],
)

h3("Model roster")
table(
    ["Role", "Model", "Used for"],
    [
        ["Primary", "llama-3.3-70b-versatile (Groq)", "Rewrite, intent+decompose, SQL, summary, recommendations, follow-ups."],
        ["Fallback 1", "openai/gpt-oss-120b", "Used automatically if the primary model fails."],
        ["Fallback 2", "qwen/qwen3-32b", "Second fallback in the chain."],
        ["Fallback 3", "meta-llama/llama-4-scout-17b-16e-instruct", "Third fallback."],
        ["Fallback 4", "llama-3.1-8b-instant", "Final, fastest fallback."],
    ],
    widths=[1.2, 3.0, 2.4],
)

h3("Comparison with alternatives")
table(
    ["Option", "Why not chosen", "Trade-off"],
    [
        ["GPT-4 (OpenAI)", "Higher cost per call; multiple LLM calls per query would be expensive.", "Very high quality but breaks the cost budget."],
        ["Self-hosted 70B", "GPU infrastructure out of budget for a student project.", "High infra cost + ops burden."],
        ["Full LangGraph rewrite", "Added ceremony without solving the real bottleneck (LLM latency).", "More abstraction, no measurable gain."],
        ["LLM-written chart code", "The model kept rewriting/mangling chart code (wrong variable, prose output).", "Replaced by deterministic Python chart specs."],
    ],
    widths=[1.6, 3.0, 2.0],
)

# ══════════════════════════════════════════════════════════════
# 4.4.8 Frontend & Backend
# ══════════════════════════════════════════════════════════════
h2("4.4.9  Frontend & Backend")

h3("Backend — FastAPI endpoints (api.py)")
para(
    "The same pipeline is exposed through a framework-free FastAPI service. Every handler "
    "installs the request’s session into the ContextVar, and the non-serialisable live Plotly "
    "figure is stripped from JSON responses (clients use chart_json)."
)
table(
    ["Method", "Path", "Description"],
    [
        ["POST", "/ask", "Run one query; return the final answer JSON."],
        ["POST", "/ask/stream", "Same, streamed as Server-Sent Events (token-by-token summary)."],
        ["POST", "/clear", "Reset a session's memory + cache."],
        ["POST", "/session/save", "Persist the session to sessions/{id}.json."],
        ["POST", "/session/load", "Restore the session from disk."],
        ["GET", "/export/csv", "Download the last result as UTF-8 CSV."],
        ["GET", "/export/pdf", "Download the full PDF report of the session."],
        ["GET", "/kpis", "Pre-computed KPI snapshot (HTML, no LLM)."],
        ["GET", "/health", "Liveness + DWH connectivity."],
        ["GET", "/  ·  /console", "Static HTML UI  ·  developer API console."],
    ],
    widths=[0.8, 1.6, 4.2],
)

h3("Frontend — static HTML UI")
para(
    "A dependency-free HTML/JS client (served at /) consumes /ask/stream and renders the answer, "
    "the result table, business recommendations, and follow-up chips. Crucially, it renders the "
    "chart from chart_json using Plotly.newPlot — the correct, portable way for any non-Gradio "
    "frontend — and offers CSV / PDF export and a KPI snapshot. A second developer console (at "
    "/console) exercises every endpoint with a one-click smoke test."
)
para("Chart rendering — the reference client integration:")
code(
    '// chart_json is a STRING → parse → Plotly.newPlot\n'
    'const spec = JSON.parse(response.chart_json);\n'
    'Plotly.newPlot("chart-div", spec.data, spec.layout, { responsive: true });'
)
para("A Gradio web app provides the same functionality for interactive/demo use; both front-ends "
     "call the identical pipeline.")

# ══════════════════════════════════════════════════════════════
# 4.4.9 Testing and Validation
# ══════════════════════════════════════════════════════════════
h2("4.4.10  Testing and Validation")
para(
    "The system was validated on three levels: a curated 200-query test set, an automated "
    "pytest suite for the deterministic core, and live end-to-end API testing."
)
h3("200-query functional test set")
para(
    "A catalogue of 200 representative queries (scripts/test_queries_200.md) spans Arabic and "
    "English across every capability: KPI, ranking, trend, distribution, comparison, compound, "
    "chart-edit, semantic cache, chitchat, and follow-up reference resolution. Running through "
    "this set drove the correctness fixes documented in Section 4.4.12."
)
h3("Automated unit tests")
para(
    "A pytest suite locks the deterministic core so refactors cannot silently regress it — the "
    "six combiners, the top-N and new-entrant logic, entity/period-label extraction, measure "
    "detection, and the chart-spec data preparation. These are pure functions, so they run fast "
    "and offline with no LLM or database calls."
)
table(
    ["Test area", "What it verifies"],
    [
        ["Combiners", "merge_on_key / pct_change / subtract / ratio / filter_by_step1 output shapes."],
        ["New-entrant guard", "A near-zero base yields NaN growth + is_new flag, never a bogus percentage."],
        ["Top-N", "Ranking results are trimmed by measure; time-series results are never row-trimmed."],
        ["Entity / period extraction", "Brand and period labels are extracted correctly for scalar comparisons (Arabic + English)."],
        ["Measure detection", "Year/month and FK/ID columns are excluded from chart measures (all-NULL SUM handled)."],
    ],
    widths=[1.8, 4.6],
)
h3("Live API validation")
para(
    "Every endpoint was exercised against a running server: /health (DB connected), /ask "
    "(valid chart_json + recommendations + follow-ups), /ask/stream (token-by-token streaming), "
    "/export/csv (UTF-8 Arabic intact), /export/pdf (valid multi-page PDF), and the session "
    "operations — all returning HTTP 200 with the expected payloads."
)

# ══════════════════════════════════════════════════════════════
# 4.4.10 Benefits
# ══════════════════════════════════════════════════════════════
h2("4.4.11  Benefits")
bullet("natural bilingual analytics with no SQL and no dashboards — an answer, a chart, an "
       "explanation, and next-step questions in seconds.", bold_prefix="Users: ")
bullet("bounded, observable cost — semantic + exact caching, 3-key rotation, transient-error "
       "retries, and per-query metrics for latency and cache-rate analysis.",
       bold_prefix="Business: ")
bullet("a reusable guarded-agent pattern — validated read-only SQL, sandboxed chart code, "
       "multi-user isolation, and a framework-free REST surface that any frontend can consume.",
       bold_prefix="Platform: ")
bullet("correct-by-construction charts and honest answers — deterministic visualisation, "
       "new-entrant guarding, and clear “no data recorded” messages instead of hallucinated "
       "numbers.", bold_prefix="Data quality: ")

# ══════════════════════════════════════════════════════════════
# 4.4.11 Integration with the Platform
# ══════════════════════════════════════════════════════════════
h2("4.4.12  Integration with the Platform")
bullet("POST /ask and /ask/stream deliver answers to the web UI (or any client); the same "
       "pipeline powers both the Gradio app and the FastAPI service.",
       bold_prefix="Frontend — chat: ")
bullet("dwh1 star schema → FAISS schema index (loaded lazily) → queried per request; the data "
       "warehouse is the single source of truth, read live on every query.",
       bold_prefix="Data flow — analytics: ")
bullet("GET /export/csv returns the last result; GET /export/pdf compiles a full session report "
       "(all queries, charts embedded as images, recommendations).", bold_prefix="Exports: ")
bullet("each query appends one JSONL record to metrics.jsonl (question, language, intent, "
       "compound flag, cache outcome, tokens, per-stage timings, total latency, success).",
       bold_prefix="Observability: ")
bullet("per-session save/load to sessions/{session_id}.json — no cross-user collisions.",
       bold_prefix="Persistence: ")
bullet("the deterministic chart theme (NAMAA indigo palette) is shared with the platform’s "
       "visual identity; the KPI snapshot endpoint feeds a live dashboard tile.",
       bold_prefix="Branding: ")

# ══════════════════════════════════════════════════════════════
# 4.4.12 Challenges and Solutions
# ══════════════════════════════════════════════════════════════
h2("4.4.13  Challenges and Solutions")
para(
    "Building a reliable Text-to-SQL analytics agent surfaced a series of real correctness "
    "problems, each resolved at the root cause rather than patched for a single query."
)
table(
    ["Challenge", "Solution"],
    [
        ["LLM-generated chart code was fragile (wrong fig variable, prose output, self-merges).",
         "Charts are built from deterministic Python specs by data shape; the LLM only picks the type."],
        ["“180,000% growth” from categories/products absent in the base year.",
         "pct_change flags a zero/negligible base as a NEW ENTRANT (NaN growth), excluded from the ranking."],
        ["Compound sub-steps used inconsistent column names (name vs product_name).",
         "filter_by_step1 matches entities by value overlap, not column name; steps trimmed to top-N first."],
        ["Compound merges produced garbage or the wrong shape.",
         "Six deterministic pandas combiners replace LLM-generated merge code; the LLM exec() path is fallback-only."],
        ["Arabic brand names transliterated to Latin → zero rows (تايد → 'Taite').",
         "Brand filters match the Arabic column; the rewrite never transliterates entity names."],
        ["discount_amount is 0 for every row in the DWH.",
         "The agent answers honestly (“no discounts recorded”) instead of retrying or hallucinating."],
        ["Exact = on Arabic names missed the definite article (منظفات vs المنظفات).",
         "Name filters use ILIKE on the core noun so article/spacing variants match."],
        ["Single-year monthly trend plotted the constant year instead of revenue.",
         "Year/month are treated as time dimensions, never as measures; trend uses the finest time column."],
        ["Free-tier token limits + per-model outages.",
         "RotatingKeyLLM: 3-key daily rotation + 4-model fallback + exponential-backoff retry on 5xx/timeout."],
        ["Concurrent users could share state.",
         "ContextVar-backed SessionState per request; asyncio propagates the context to all child tasks."],
    ],
    widths=[3.2, 3.4],
)

# ══════════════════════════════════════════════════════════════
# Summary
# ══════════════════════════════════════════════════════════════
h2("Summary")
para("The NAMAA Analytics Agent is critical to the platform because it:")
bullet("Turns a live retail data warehouse into a natural bilingual conversation — no dashboards, no SQL.")
bullet("Generates, validates, and executes read-only SQL autonomously, with multi-layer safety and self-correction.")
bullet("Visualises results deterministically and explains them, with recommendations and follow-up questions.")
bullet("Answers honestly — “not available” / “no discounts recorded” instead of fabricated figures.")
bullet("Operates safely and measurably — sandboxed code, per-session isolation, caching, key rotation, and per-query metrics — behind both a Gradio UI and a framework-free FastAPI endpoint.")

# ── save ──────────────────────────────────────────────────────
OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "NAMAA_Analytics_Agent.docx")
doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
